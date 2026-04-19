# -*- coding: utf-8 -*-
# 更正了所有问题,目前就只剩下了那个各级标题的字体不对了
import re
from docx import Document
from openpyxl import load_workbook
from docx.oxml.ns import qn  # 导入 qn 函数
from docx.oxml import OxmlElement  # 导入 OxmlElement 用于创建 XML 元素


def replace_placeholders_in_paragraph(paragraph, wb, pattern):
    """
    替换段落中的占位符，保留格式和样式，同时处理跨 run 的占位符。
    """
    # 记录原始 runs 的文本和样式信息
    run_texts = []
    run_styles = []
    for run in paragraph.runs:
        run_texts.append(run.text)
        style = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'size': run.font.size,
            'color': run.font.color.rgb if run.font.color.rgb else None,
            'name': run.font.name,
            'strike': run.font.strike,
            'subscript': run.font.subscript,
            'superscript': run.font.superscript,
        }
        run_styles.append(style)

    # 合并文本
    full_text = ''.join(run_texts)
    if not full_text:
        return

    # 查找占位符
    matches = list(pattern.finditer(full_text))
    if not matches:
        return  # 无占位符，返回

    # 构建字符索引到样式的映射
    char_styles = []
    for text, style in zip(run_texts, run_styles):
        char_styles.extend([style] * len(text))

    # 开始替换占位符
    new_text = ''
    new_styles = []
    last_index = 0
    for match in matches:
        sheet_name, cell_address = match.groups()
        start, end = match.start(), match.end()

        # 添加占位符前的文本和样式
        new_text += full_text[last_index:start]
        new_styles.extend(char_styles[last_index:start])

        # 获取替换值
        try:
            value = wb[sheet_name][cell_address].value
            if isinstance(value, (int, float)):
                value = f"{value:.2f}"  # 数值保留两位小数
            elif value is None:
                value = ''
        except KeyError:
            value = ''
        value = str(value)

        # 获取占位符的样式（取占位符开始位置的样式）
        if start < len(char_styles):
            placeholder_style = char_styles[start]
        else:
            # 如果找不到占位符的样式，使用默认样式
            placeholder_style = {
                'bold': False,
                'italic': False,
                'underline': False,
                'size': None,
                'color': None,
                'name': None,
                'strike': False,
                'subscript': False,
                'superscript': False,
            }

        # 添加替换值和样式
        new_text += value
        new_styles.extend([placeholder_style] * len(value))

        last_index = end

    # 添加剩余的文本和样式
    new_text += full_text[last_index:]
    new_styles.extend(char_styles[last_index:])

    # 清空原段落
    p = paragraph._element
    p.clear_content()

    # 重新创建 runs
    current_style = None
    current_run = None
    for idx, char in enumerate(new_text):
        style = new_styles[idx]
        # 将样式字典转换为元组以进行比较
        style_tuple = tuple(sorted(style.items()))
        if style_tuple != current_style:
            current_run = paragraph.add_run()
            font = current_run.font
            current_run.bold = style['bold']
            current_run.italic = style['italic']
            current_run.underline = style['underline']
            font.size = style['size']
            if style['color']:
                font.color.rgb = style['color']
            else:
                font.color.rgb = None  # 如果颜色为 None，则清除颜色
            # 设置字体名称，应用到所有字体属性
            if style['name']:
                font.name = style['name']
                # 确保 rPr 和 rFonts 存在
                if font.element.rPr is None:
                    font.element._element.append(OxmlElement('w:rPr'))
                rPr = font.element.rPr
                if rPr.rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                else:
                    rFonts = rPr.rFonts
                rFonts.set(qn('w:eastAsia'), style['name'])
                rFonts.set(qn('w:ascii'), style['name'])
                rFonts.set(qn('w:hAnsi'), style['name'])
                rFonts.set(qn('w:cs'), style['name'])
            else:
                # 如果没有指定字体名称，且 rFonts 存在，则移除它
                if font.element.rPr is not None and font.element.rPr.rFonts is not None:
                    font.element.rPr.remove(font.element.rPr.rFonts)
            font.strike = style['strike']
            font.subscript = style['subscript']
            font.superscript = style['superscript']
            current_style = style_tuple
        current_run.add_text(char)


def replace_placeholders_in_table(table, wb, pattern):
    """
    替换表格中的占位符。
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, wb, pattern)


def replace_placeholders_in_header_footer(header_footer, wb, pattern):
    """
    替换页眉或页脚中的占位符。
    """
    for paragraph in header_footer.paragraphs:
        replace_placeholders_in_paragraph(paragraph, wb, pattern)
    for table in header_footer.tables:
        replace_placeholders_in_table(table, wb, pattern)


def replace_placeholders_in_shapes(shapes, wb, pattern):
    """
    替换文本框中的占位符。
    """
    for shape in shapes:
        if shape.has_text_frame:
            for paragraph in shape.text_frame.paragraphs:
                replace_placeholders_in_paragraph(paragraph, wb, pattern)


def replace_placeholders(word_file, excel_file, output_file):
    """
    主函数：加载 Word 和 Excel 文件，替换占位符并保存。
    """
    # 加载 Word 文档和 Excel 工作簿
    doc = Document(word_file)
    wb = load_workbook(excel_file, data_only=True)

    # 定义占位符正则表达式 {SheetName!CellAddress}
    pattern = re.compile(r'\{([^!}]+)!([A-Z]+[0-9]+)\}')

    # 替换正文段落中的占位符
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, wb, pattern)

    # 替换表格中的占位符
    for table in doc.tables:
        replace_placeholders_in_table(table, wb, pattern)

    # 替换页眉和页脚中的占位符
    for section in doc.sections:
        replace_placeholders_in_header_footer(section.header, wb, pattern)
        replace_placeholders_in_header_footer(section.footer, wb, pattern)

    # 替换文本框中的占位符
    if hasattr(doc, 'shapes'):
        replace_placeholders_in_shapes(doc.shapes, wb, pattern)

    # 保存修改后的文档
    doc.save(output_file)